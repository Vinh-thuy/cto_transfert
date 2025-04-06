from docx import Document
import json

def extraire_contenu_word(chemin_fichier):
    doc = Document(chemin_fichier)
    contenu = {"paragraphes": [], "tableaux": []}

    # Extraire les paragraphes
    for para in doc.paragraphs:
        contenu["paragraphes"].append(para.text)

    # Extraire les tableaux
    for table in doc.tables:
        tableau = []
        for ligne in table.rows:
            ligne_tableau = [cellule.text for cellule in ligne.cells]
            tableau.append(ligne_tableau)
        contenu["tableaux"].append(tableau)

    return contenu

chemin_fichier = 'votre_document.docx'
contenu = extraire_contenu_word(chemin_fichier)

# Convertir le contenu en JSON
contenu_json = json.dumps(contenu, ensure_ascii=False, indent=4)

# Sauvegarder le JSON dans un fichier
with open('contenu_document.json', 'w', encoding='utf-8') as fichier_json:
    fichier_json.write(contenu_json)
